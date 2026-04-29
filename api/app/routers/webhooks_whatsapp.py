"""Inbound WhatsApp webhook (Gupshup v2 payload format).

Reporters forward government press releases to the Vrittant WABA number
(+91 89843 36534, hosted by Gupshup). For each inbound message:

1. Look up the sender phone in the `users` table.
2. If registered + active, create a draft `Story` assigned to them with
   `source="whatsapp"` and the message text as the body. Reply with a
   confirmation.
3. If not registered, reply with a polite "contact your editor" note and
   drop the message.

Always returns 200 — Gupshup retries on any non-2xx, which would amplify
errors and double-create stories. Idempotency is enforced via the
`whatsapp_inbound_dedup` table (PK on Gupshup's `message.id`).
"""
import json
import logging
import os
import uuid
from datetime import timedelta

import httpx
from fastapi import APIRouter, Depends, Request
from sqlalchemy.orm import Session

from ..database import get_db
from ..models.story import Story
from ..models.user import User
from ..models.webhook_dedup import WhatsappInboundDedup
from ..services import sarvam_client, storage
from ..services import whatsapp_classifier as classifier
from ..services.assignment import NoReviewersAvailable, pick_assignee
from ..services.categorizer import classify_category, org_category_keys
from ..utils.tz import now_ist

router = APIRouter(prefix="/webhooks/whatsapp", tags=["webhooks"])
logger = logging.getLogger(__name__)

STITCH_MINUTES = int(os.environ.get("WHATSAPP_STITCH_MINUTES", "10"))
CLOSE_KEYWORDS = {"done", "/done", "end", "/end"}


def _is_close_keyword(text: str) -> bool:
    return text.strip().lower() in CLOSE_KEYWORDS


# org_category_keys is now shared with the editor save path; see
# services/categorizer.py. The legacy local helper used to live here.


def _find_open_draft(db, user_id: str):
    return (
        db.query(Story)
        .filter(
            Story.reporter_id == user_id,
            Story.source == "whatsapp",
            Story.whatsapp_session_open_until.isnot(None),
            Story.whatsapp_session_open_until > now_ist(),
        )
        .order_by(Story.whatsapp_session_open_until.desc())
        .first()
    )

# The Vrittant WABA number provisioned in Gupshup. Hard-coded because it
# identifies the source of *outbound* replies; if the number ever changes
# this is the one place to update.
GUPSHUP_SOURCE_PHONE = "918984336534"
GUPSHUP_APP_NAME = "Vrittant"
GUPSHUP_API_URL = "https://api.gupshup.io/wa/api/v1/msg"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _normalize_phone(raw: str) -> str:
    """Gupshup sends '919...' (no plus); our DB stores '+919...'. Normalize."""
    raw = (raw or "").strip()
    if not raw:
        return raw
    return raw if raw.startswith("+") else "+" + raw


def _extract_content(inner_type: str, inner_payload: dict) -> tuple[str, str | None]:
    """Return (text, media_url) from a Gupshup v2 inner payload.

    Gupshup wraps the actual content under `payload.payload` and the type
    under `payload.type`. For text it's `{"text": "..."}`; for media it's
    `{"url": "...", "caption": "...", "name": "...", "contentType": "..."}`.
    """
    if inner_type == "text":
        return (inner_payload.get("text") or "", None)
    if inner_type in ("image", "video", "audio", "document", "file"):
        caption = inner_payload.get("caption") or inner_payload.get("name") or ""
        return (caption, inner_payload.get("url"))
    return ("", None)


def _media_type_for(inner_type: str) -> str:
    """Map Gupshup inbound type → the `media_type` value the panel groups by.

    The reviewer panel filters paragraphs into image / audio / generic media
    rails based on this field (see helpers.js → transformStory)."""
    if inner_type == "image":
        return "photo"
    if inner_type in ("video", "audio"):
        return inner_type
    return "file"


# Map common Gupshup contentType values → file extension for the GCS object.
# Office formats matter for two reasons: (a) without an extension the
# downloaded file lands as `<hex>` and most OSes refuse to open it; (b) the
# panel's attachment rail uses the extension to pick the right icon/label.
_CONTENT_TYPE_EXT = {
    "image/jpeg": ".jpg", "image/jpg": ".jpg", "image/png": ".png",
    "image/webp": ".webp", "image/gif": ".gif", "image/heic": ".heic",
    "video/mp4": ".mp4", "video/quicktime": ".mov", "video/webm": ".webm",
    "audio/ogg": ".ogg", "audio/mpeg": ".mp3", "audio/mp4": ".m4a",
    "audio/aac": ".aac", "audio/amr": ".amr",
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.ms-excel": ".xls",
    "text/plain": ".txt",
    "text/csv": ".csv",
}


def _extract_docx_text(body: bytes) -> str:
    """Pull the visible paragraph text out of a .docx blob. Returns "" on
    any parse error — docx extraction is a best-effort enhancement; the
    file itself is still attached to the story regardless."""
    try:
        import io
        from docx import Document  # python-docx
        doc = Document(io.BytesIO(body))
        chunks: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                chunks.append(t)
        # Tables are common in government press releases — flatten them
        # row-by-row, tab-separating cells, so the reviewer at least sees
        # the data even if the layout doesn't survive.
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = "\t".join(c for c in cells if c)
                if line:
                    chunks.append(line)
        return "\n\n".join(chunks).strip()
    except Exception:
        logger.exception("Failed to extract docx text")
        return ""


async def _persist_media(
    url: str, content_type: str | None = None, original_name: str | None = None
) -> tuple[str | None, bytes | None, str | None]:
    """Download a media file from Gupshup's transient URL and re-host it
    in our own bucket.

    Returns ``(stored_url, body, content_type)``. ``stored_url`` is None
    only if the upload failed (callers fall back to the Gupshup link).
    ``body`` and ``content_type`` are returned even when upload fails, so
    the caller can still extract docx text from the bytes we did fetch.
    """
    try:
        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            body = resp.content
            # Prefer the response's Content-Type; fall back to Gupshup's hint.
            ct = (resp.headers.get("content-type") or content_type or "").split(";")[0].strip()
    except Exception:
        logger.exception("Failed to download Gupshup media from %s", url)
        return (None, None, None)

    ext = _CONTENT_TYPE_EXT.get(ct.lower(), "")
    if not ext and original_name and "." in original_name:
        # Reporter's original filename usually carries the right extension
        # even when Gupshup's contentType is generic application/octet-stream.
        ext = "." + original_name.rsplit(".", 1)[-1].lower()
    if not ext:
        # Try the URL path as a last resort (e.g. /abc.jpg).
        url_path = url.split("?", 1)[0]
        if "." in url_path.rsplit("/", 1)[-1]:
            ext = "." + url_path.rsplit(".", 1)[-1].lower()
    filename = f"{uuid.uuid4().hex}{ext}"

    try:
        stored = storage.save_file(body, filename, subfolder="whatsapp")
    except Exception:
        logger.exception("Failed to upload WhatsApp media to storage")
        stored = None
    return (stored, body, ct)


def _is_docx_ct(ct: str | None) -> bool:
    return (ct or "").lower() == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def _send_gupshup_reply(to_phone: str, text: str) -> None:
    """POST a text reply via Gupshup. No-op when no API key is configured
    (e.g. tests, local dev) so the integration degrades gracefully."""
    api_key = os.environ.get("GUPSHUP_API_KEY", "")
    if not api_key:
        logger.info("GUPSHUP_API_KEY unset; would have replied to %s: %s", to_phone, text)
        return
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            await client.post(
                GUPSHUP_API_URL,
                headers={"apikey": api_key},
                data={
                    "channel": "whatsapp",
                    "source": GUPSHUP_SOURCE_PHONE,
                    "destination": to_phone,
                    "message": json.dumps({"type": "text", "text": text}),
                    "src.name": GUPSHUP_APP_NAME,
                },
            )
    except Exception:  # pragma: no cover — best-effort reply
        logger.exception("Gupshup reply to %s failed", to_phone)


# ---------------------------------------------------------------------------
# Webhook
# ---------------------------------------------------------------------------
@router.post("/gupshup")
async def gupshup_inbound(request: Request, db: Session = Depends(get_db)):
    body = await request.json()

    # Gupshup posts many event types; we only care about user-sent messages.
    # Status callbacks (sent/delivered/read/failed) come as type="message-event".
    if body.get("type") != "message":
        return {"ok": True, "skipped": "non-message event"}

    payload = body.get("payload") or {}
    msg_id = payload.get("id")
    sender_raw = (payload.get("sender") or {}).get("phone") or payload.get("source")
    if not msg_id or not sender_raw:
        return {"ok": True, "skipped": "missing id or sender"}

    # Idempotency check first — duplicate retries should be the cheapest path.
    if db.query(WhatsappInboundDedup).filter_by(message_id=msg_id).first():
        return {"ok": True, "skipped": "duplicate"}

    sender_phone = _normalize_phone(sender_raw)
    user = (
        db.query(User)
        .filter(User.phone == sender_phone, User.is_active == True)  # noqa: E712
        .first()
    )

    # Record the dedup row regardless of outcome — we don't want to keep
    # re-processing an unknown sender on every Gupshup retry.
    db.add(WhatsappInboundDedup(message_id=msg_id, received_at=now_ist()))

    if not user:
        db.commit()
        await _send_gupshup_reply(
            sender_raw,
            "This number is not registered with Vrittant. Please contact your editor.",
        )
        return {"ok": True, "skipped": "unknown sender"}

    inner_type = payload.get("type", "text")
    inner_payload = payload.get("payload") or {}
    text, media_url = _extract_content(inner_type, inner_payload)

    # 1. Close keyword (only when no media — a photo with caption "done"
    #    should append, not seal).
    if text and not media_url and _is_close_keyword(text):
        open_draft = _find_open_draft(db, user.id)
        if open_draft is not None:
            open_draft.whatsapp_session_open_until = None
            n = len(open_draft.paragraphs or [])
            db.commit()
            await _send_gupshup_reply(
                sender_raw,
                f"Story #{open_draft.id[:8]} sealed with {n} items.",
            )
            return {"ok": True, "closed": open_draft.id}
        db.commit()
        await _send_gupshup_reply(sender_raw, "No open story to close.")
        return {"ok": True, "skipped": "close-no-draft"}

    # 2. Open draft → append.
    open_draft = _find_open_draft(db, user.id)
    if open_draft is not None:
        stored_media_url: str | None = None
        media_body: bytes | None = None
        media_ct: str | None = None
        original_name = inner_payload.get("name") or None
        if media_url:
            stored_media_url, media_body, media_ct = await _persist_media(
                media_url, inner_payload.get("contentType"), original_name
            )
            stored_media_url = stored_media_url or media_url

        paragraphs = list(open_draft.paragraphs or [])
        if text:
            paragraphs.append({"id": str(uuid.uuid4()), "text": text})
        if stored_media_url:
            # Empty `text` so the editor body doesn't show the URL inline —
            # the panel reads `media_path` to build the Attachments rail.
            media_para = {
                "id": str(uuid.uuid4()),
                "text": "",
                "media_path": stored_media_url,
                "media_type": _media_type_for(inner_type),
            }
            if original_name:
                # `media_name` drives the panel's display name + download
                # filename. Without this, docs land as `<hex>` with no
                # extension on disk and the editor falls back to the URL
                # basename (also extensionless).
                media_para["media_name"] = original_name
            paragraphs.append(media_para)
        # Best-effort docx text extraction. We append it as a separate
        # text paragraph after the file attachment so the reviewer sees
        # the body content inline without losing the original document.
        if media_body and _is_docx_ct(media_ct):
            extracted = _extract_docx_text(media_body)
            if extracted:
                paragraphs.append({"id": str(uuid.uuid4()), "text": extracted})
        open_draft.paragraphs = paragraphs
        open_draft.whatsapp_session_open_until = now_ist() + timedelta(minutes=STITCH_MINUTES)
        open_draft.refresh_search_text()
        db.commit()
        return {"ok": True, "appended_to": open_draft.id}

    # 3 / 4. No open draft — maybe classify, then create new story.
    # Media-only messages skip the classifier (always news intent).
    needs_triage = False
    if text:
        # Classification happens before we have a story_id — bucket it
        # so we can see how much classifier overhead the WhatsApp intake
        # path is costing us per month.
        with sarvam_client.cost_context(bucket="whatsapp_intake"):
            label = await classifier.classify(text)
        if label == "chitchat":
            db.commit()
            await _send_gupshup_reply(
                sender_raw,
                "Got your message. To submit news, forward press releases or news text here.",
            )
            return {"ok": True, "skipped": "chitchat"}
        needs_triage = (label == "unclear")

    stored_media_url = None
    media_body: bytes | None = None
    media_ct: str | None = None
    original_name = inner_payload.get("name") or None
    if media_url:
        stored_media_url, media_body, media_ct = await _persist_media(
            media_url, inner_payload.get("contentType"), original_name
        )
        stored_media_url = stored_media_url or media_url

    # Headline = the first non-empty line of the forwarded text, capped
    # at 120 chars to fit the editorial display width. Falls back to
    # "Forwarded from WhatsApp" when the message has no text body
    # (image-only forwards). Reviewers reference these stories by the
    # display_id ("PNS-26-1234") which is also unique per-org and
    # human-readable, replacing the older "#shortid" prefix that used
    # to live on the headline itself.
    new_story_id = str(uuid.uuid4())
    if text:
        first_line = text.split("\n", 1)[0].strip()
    else:
        first_line = ""
    if first_line:
        if len(first_line) > 120:
            first_line = first_line[:119].rstrip() + "…"
        headline = first_line
    else:
        headline = "Forwarded from WhatsApp"
    paragraphs = []
    if text:
        paragraphs.append({"id": str(uuid.uuid4()), "text": text})
    if stored_media_url:
        media_para = {
            "id": str(uuid.uuid4()),
            "text": "",
            "media_path": stored_media_url,
            "media_type": _media_type_for(inner_type),
        }
        if original_name:
            media_para["media_name"] = original_name
        paragraphs.append(media_para)
    # Best-effort docx text extraction (mirrors the open-draft branch).
    if media_body and _is_docx_ct(media_ct):
        extracted = _extract_docx_text(media_body)
        if extracted:
            paragraphs.append({"id": str(uuid.uuid4()), "text": extracted})
            # Use extracted text for category classification too — a docx
            # forward with no caption would otherwise be Uncategorized.
            if not text:
                text = extracted

    # Best-effort category tag at ingestion. We pull the org's configured
    # category keys (falls back to the platform defaults if the org has
    # no override). If Sarvam errors out the helper returns None and we
    # leave `category` unset — same as the pre-categorisation behaviour.
    category_keys = org_category_keys(db, user.organization_id)
    category = await classify_category(text, category_keys) if text else None

    from ..services.story_seq import assign_next_seq
    story = Story(
        id=new_story_id,
        organization_id=user.organization_id,
        seq_no=assign_next_seq(db, user.organization_id),
        reporter_id=user.id,
        # `assigned_to` is filled below via pick_assignee for reporters so
        # the story enters the same review queue as panel-created stories.
        # Reviewers/admins keep self-assign — they're working on it.
        assigned_to=None if user.user_type == "reporter" else user.id,
        assigned_match_reason=None if user.user_type == "reporter" else "manual",
        headline=headline,
        category=category,
        paragraphs=paragraphs,
        status="submitted",
        submitted_at=now_ist(),
        source="whatsapp",
        whatsapp_session_open_until=now_ist() + timedelta(minutes=STITCH_MINUTES),
        needs_triage=needs_triage,
    )
    if user.user_type == "reporter":
        try:
            reviewer, reason = pick_assignee(story, db)
            story.assigned_to = reviewer.id
            story.assigned_match_reason = reason
        except NoReviewersAvailable:
            # Org has no active reviewers — leave unassigned so it surfaces
            # in the unassigned queue rather than stuck on the reporter.
            logger.warning(
                "WA inbound from reporter %s: no reviewers in org %s, story %s left unassigned",
                user.id, user.organization_id, story.id,
            )
    story.refresh_search_text()
    db.add(story)
    db.commit()

    await _send_gupshup_reply(
        sender_raw,
        f"Started story #{story.id[:8]}. Send more text or media; reply done or wait {STITCH_MINUTES} min to close.",
    )
    return {"ok": True, "story_id": story.id}
